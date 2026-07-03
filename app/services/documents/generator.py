"""
app/services/documents/generator.py
-------------------------------------
Converts Claude's structured JSON output into downloadable documents.

Supported output formats:
  - PDF  — WeasyPrint renders HTML templates to professional PDFs
  - DOCX — python-docx for editable Word documents
  - TEXT — Plain formatted text for in-chat Telegram preview (Free tier)

All templates use Jinja2 for HTML rendering before PDF conversion.
"""

from __future__ import annotations

import io
import os
from datetime import datetime
from pathlib import Path
from typing import Optional

from jinja2 import Environment, DictLoader
from loguru import logger

from app.core.constants import DocType, OutputFormat, SubscriptionTier

# Determine WeasyPrint availability
try:
    from weasyprint import HTML as WeasyHTML
    WEASYPRINT_AVAILABLE = True
except (ImportError, OSError):
    WEASYPRINT_AVAILABLE = False
    logger.warning("WeasyPrint not available — PDF generation will be skipped")

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# ── HTML Templates ────────────────────────────────────────────────────────────

INVOICE_HTML = """
<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8">
<style>
  body { font-family: Arial, sans-serif; font-size: 11pt; color: #1a1a1a; margin: 0; padding: 0; }
  .page { padding: 40px 50px; max-width: 750px; margin: 0 auto; }
  .header { display: flex; justify-content: space-between; align-items: flex-start; margin-bottom: 30px; border-bottom: 3px solid #1B5E20; padding-bottom: 20px; }
  .brand { font-size: 22pt; font-weight: bold; color: #1B5E20; }
  .brand-sub { font-size: 9pt; color: #555; }
  .invoice-meta { text-align: right; }
  .invoice-title { font-size: 18pt; color: #1B5E20; font-weight: bold; }
  .invoice-num { font-size: 10pt; color: #333; margin-top: 5px; }
  .parties { display: flex; justify-content: space-between; margin-bottom: 25px; }
  .party-box { width: 48%; }
  .party-label { font-size: 9pt; color: #888; text-transform: uppercase; letter-spacing: 1px; margin-bottom: 5px; }
  .party-name { font-size: 12pt; font-weight: bold; color: #1B5E20; }
  .party-detail { font-size: 9.5pt; color: #444; line-height: 1.5; }
  table { width: 100%; border-collapse: collapse; margin-bottom: 20px; }
  thead tr { background-color: #1B5E20; color: white; }
  thead th { padding: 10px 12px; text-align: left; font-size: 10pt; }
  tbody tr:nth-child(even) { background-color: #f9f9f9; }
  tbody td { padding: 9px 12px; font-size: 10pt; border-bottom: 1px solid #e8e8e8; }
  .text-right { text-align: right; }
  .totals { margin-left: auto; width: 280px; }
  .totals-row { display: flex; justify-content: space-between; padding: 5px 0; font-size: 10pt; }
  .totals-total { display: flex; justify-content: space-between; padding: 10px 0; font-size: 13pt; font-weight: bold; border-top: 2px solid #1B5E20; color: #1B5E20; }
  .bank-box { background: #f0f7f0; border-left: 4px solid #1B5E20; padding: 15px; margin-top: 20px; }
  .bank-title { font-weight: bold; font-size: 10pt; color: #1B5E20; margin-bottom: 8px; }
  .bank-detail { font-size: 10pt; color: #333; line-height: 1.6; }
  .footer { text-align: center; margin-top: 30px; font-size: 9pt; color: #888; border-top: 1px solid #e8e8e8; padding-top: 15px; }
  {% if watermark %}.watermark { position: fixed; top: 50%; left: 50%; transform: translate(-50%,-50%) rotate(-45deg); font-size: 60pt; color: rgba(200,200,200,0.25); z-index: -1; white-space: nowrap; }{% endif %}
</style>
</head>
<body>
{% if watermark %}<div class="watermark">BIZPILOT NG FREE</div>{% endif %}
<div class="page">
  <div class="header">
    <div>
      <div class="brand">{{ data.seller.name }}</div>
      <div class="brand-sub">{% if data.seller.cac %}RC: {{ data.seller.cac }}{% endif %}{% if data.seller.tin %} | TIN: {{ data.seller.tin }}{% endif %}</div>
    </div>
    <div class="invoice-meta">
      <div class="invoice-title">INVOICE</div>
      <div class="invoice-num">{{ data.invoice_number }}</div>
      <div style="font-size:9.5pt;color:#555;margin-top:5px;">Issue: {{ data.issue_date }} | Due: {{ data.due_date }}</div>
    </div>
  </div>

  <div class="parties">
    <div class="party-box">
      <div class="party-label">Bill From</div>
      <div class="party-name">{{ data.seller.name }}</div>
      <div class="party-detail">{{ data.seller.address }}</div>
    </div>
    <div class="party-box">
      <div class="party-label">Bill To</div>
      <div class="party-name">{{ data.bill_to.name }}</div>
      <div class="party-detail">{{ data.bill_to.email }}<br>{{ data.bill_to.address }}</div>
    </div>
  </div>

  <table>
    <thead>
      <tr>
        <th>#</th><th>Description</th><th>Qty</th><th class="text-right">Unit Price</th><th class="text-right">Total</th>
      </tr>
    </thead>
    <tbody>
      {% for item in data.items %}
      <tr>
        <td>{{ loop.index }}</td>
        <td>{{ item.description }}</td>
        <td>{{ item.quantity }}</td>
        <td class="text-right">₦{{ "{:,.2f}".format(item.unit_price) }}</td>
        <td class="text-right">₦{{ "{:,.2f}".format(item.line_total) }}</td>
      </tr>
      {% endfor %}
    </tbody>
  </table>

  <div class="totals">
    <div class="totals-row"><span>Subtotal</span><span>₦{{ "{:,.2f}".format(data.subtotal) }}</span></div>
    <div class="totals-row"><span>VAT ({{ data.vat_rate }}%)</span><span>₦{{ "{:,.2f}".format(data.vat_amount) }}</span></div>
    {% if data.wht_amount > 0 %}<div class="totals-row" style="color:#c00"><span>WHT ({{ data.wht_rate }}%) deducted</span><span>-₦{{ "{:,.2f}".format(data.wht_amount) }}</span></div>{% endif %}
    <div class="totals-total"><span>TOTAL DUE</span><span>₦{{ "{:,.2f}".format(data.total_payable) }}</span></div>
  </div>

  <div class="bank-box">
    <div class="bank-title">Payment Details — {{ data.payment_terms }}</div>
    <div class="bank-detail">
      Account Name: <strong>{{ data.bank_details.account_name }}</strong><br>
      Account Number: <strong>{{ data.bank_details.account_number }}</strong><br>
      Bank: <strong>{{ data.bank_details.bank_name }}</strong>
    </div>
  </div>

  {% if data.notes %}<p style="font-size:9.5pt;color:#555;margin-top:15px;"><strong>Notes:</strong> {{ data.notes }}</p>{% endif %}

  <div class="footer">
    {{ data.footer }}<br>
    <strong>Generated by BizPilot NG</strong> | bizpilot.ng
  </div>
</div>
</body>
</html>
"""

PROPOSAL_HTML = """
<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8">
<style>
  body { font-family: Georgia, serif; font-size: 11pt; color: #1a1a1a; margin: 0; }
  .page { padding: 45px 55px; max-width: 750px; margin: 0 auto; }
  .header { border-bottom: 3px solid #1A237E; padding-bottom: 20px; margin-bottom: 25px; }
  h1 { color: #1A237E; font-size: 20pt; margin: 0 0 5px; }
  h2 { color: #1A237E; font-size: 13pt; border-bottom: 1px solid #c5cae9; padding-bottom: 5px; margin-top: 22px; }
  .meta { font-size: 9.5pt; color: #555; }
  .section { margin-bottom: 18px; line-height: 1.7; }
  .amount-box { background: #e8eaf6; border-left: 5px solid #1A237E; padding: 15px 20px; margin: 20px 0; }
  .amount-total { font-size: 16pt; font-weight: bold; color: #1A237E; }
  table { width: 100%; border-collapse: collapse; margin: 10px 0; }
  th { background: #1A237E; color: white; padding: 8px 10px; text-align: left; font-size: 10pt; }
  td { padding: 8px 10px; border-bottom: 1px solid #e8e8e8; font-size: 10pt; }
  .signature-block { display: flex; justify-content: space-between; margin-top: 40px; }
  .sig-box { width: 45%; border-top: 1px solid #333; padding-top: 8px; font-size: 9.5pt; }
  .footer { text-align: center; margin-top: 30px; font-size: 9pt; color: #888; border-top: 1px solid #e8e8e8; padding-top: 15px; }
  {% if watermark %}.watermark { position: fixed; top: 50%; left: 50%; transform: translate(-50%,-50%) rotate(-45deg); font-size: 60pt; color: rgba(200,200,200,0.25); z-index: -1; white-space: nowrap; }{% endif %}
</style>
</head>
<body>
{% if watermark %}<div class="watermark">BIZPILOT NG FREE</div>{% endif %}
<div class="page">
  <div class="header">
    <div class="meta">{{ data.date }} | Ref: {{ data.proposal_number }} | Valid until {{ data.valid_until }}</div>
    <h1>Business Proposal</h1>
    <div style="font-size:10.5pt"><strong>Prepared for:</strong> {{ data.to.name }}{% if data.to.company %}, {{ data.to.company }}{% endif %}</div>
    <div style="font-size:10.5pt"><strong>Prepared by:</strong> {{ data.from.name }}, {{ data.from.business }}</div>
  </div>

  <h2>Executive Summary</h2>
  <div class="section">{{ data.executive_summary }}</div>

  <h2>Project Overview</h2>
  <div class="section">{{ data.project_overview }}</div>

  <h2>Scope of Work</h2>
  <div class="section">{% if data.scope_of_work is string %}{{ data.scope_of_work }}{% else %}<ul>{% for s in data.scope_of_work %}<li>{{ s }}</li>{% endfor %}</ul>{% endif %}</div>

  <h2>Deliverables</h2>
  <div class="section"><ul>{% for d in data.deliverables %}<li>{% if d is string %}{{ d }}{% else %}{{ d.item | default(d) }}{% if d.timeline is defined %} — {{ d.timeline }}{% endif %}{% endif %}</li>{% endfor %}</ul></div>

  <h2>Timeline</h2>
  <div class="section">{{ data.timeline }}</div>

  <h2>Investment</h2>
  <div class="amount-box">
    <div>Total Project Value</div>
    <div class="amount-total">₦{{ "{:,.2f}".format(data.investment.total_amount) }}</div>
  </div>
  {% if data.investment.breakdown %}
  <table>
    <tr><th>Milestone</th><th>Percentage</th><th>Amount</th></tr>
    {% for m in data.investment.breakdown %}
    <tr><td>{{ m.milestone }}</td><td>{{ m.percentage }}%</td><td>₦{{ "{:,.2f}".format(m.amount) }}</td></tr>
    {% endfor %}
  </table>
  {% endif %}

  <h2>Why Choose Us</h2>
  <div class="section">{{ data.why_choose_us }}</div>

  <h2>Terms & Conditions</h2>
  <div class="section" style="font-size:9.5pt;color:#444">{{ data.terms_and_conditions }}</div>

  <div class="section">{{ data.closing_message }}</div>

  <div class="signature-block">
    <div class="sig-box">
      <strong>{{ data.from.name }}</strong><br>{{ data.from.business }}<br>Date: ________________
    </div>
    <div class="sig-box">
      <strong>{{ data.to.name }}</strong><br>{{ data.to.company }}<br>Date: ________________
    </div>
  </div>

  <div class="footer">Generated by <strong>BizPilot NG</strong> | bizpilot.ng</div>
</div>
</body>
</html>
"""

CONTRACT_HTML = """
<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8">
<style>
  body { font-family: Georgia, serif; font-size: 10.5pt; color: #1a1a1a; margin: 0; }
  .page { padding: 45px 55px; max-width: 750px; margin: 0 auto; }
  .header { text-align: center; border-bottom: 3px solid #263238; padding-bottom: 20px; margin-bottom: 25px; }
  h1 { color: #263238; font-size: 18pt; margin: 0 0 5px; text-transform: uppercase; letter-spacing: 2px; }
  .contract-ref { font-size: 9.5pt; color: #555; }
  h2 { color: #263238; font-size: 12pt; border-bottom: 1px solid #b0bec5; padding-bottom: 4px; margin-top: 22px; text-transform: uppercase; letter-spacing: 1px; }
  .section { margin-bottom: 15px; line-height: 1.8; text-align: justify; }
  .parties { display: flex; justify-content: space-between; margin: 20px 0; }
  .party-box { width: 46%; background: #eceff1; padding: 15px; border-left: 4px solid #263238; }
  .party-label { font-size: 9pt; color: #888; text-transform: uppercase; letter-spacing: 1px; margin-bottom: 5px; }
  .party-name { font-size: 12pt; font-weight: bold; color: #263238; }
  .party-detail { font-size: 9.5pt; color: #444; line-height: 1.5; }
  .def-term { font-weight: bold; }
  .signature-block { display: flex; justify-content: space-between; margin-top: 50px; padding-top: 10px; }
  .sig-box { width: 44%; }
  .sig-line { border-top: 1px solid #333; margin-top: 60px; padding-top: 8px; font-size: 9.5pt; }
  .seal-space { border: 1px dashed #aaa; width: 80px; height: 80px; margin-top: 10px; display: flex; align-items: center; justify-content: center; font-size: 8pt; color: #aaa; }
  .footer { text-align: center; margin-top: 30px; font-size: 9pt; color: #888; border-top: 1px solid #e8e8e8; padding-top: 15px; }
  {% if watermark %}.watermark { position: fixed; top: 50%; left: 50%; transform: translate(-50%,-50%) rotate(-45deg); font-size: 60pt; color: rgba(200,200,200,0.25); z-index: -1; white-space: nowrap; }{% endif %}
</style>
</head>
<body>
{% if watermark %}<div class="watermark">BIZPILOT NG FREE</div>{% endif %}
<div class="page">
  <div class="header">
    <h1>{{ data.contract_title }}</h1>
    <div class="contract-ref">Ref: {{ data.contract_number }} | Date: {{ data.date }}</div>
  </div>

  <div class="parties">
    <div class="party-box">
      <div class="party-label">Party A — {{ data.parties.party_a.role }}</div>
      <div class="party-name">{{ data.parties.party_a.name }}</div>
      <div class="party-detail">{{ data.parties.party_a.address }}</div>
    </div>
    <div class="party-box">
      <div class="party-label">Party B — {{ data.parties.party_b.role }}</div>
      <div class="party-name">{{ data.parties.party_b.name }}</div>
      <div class="party-detail">{{ data.parties.party_b.address }}</div>
    </div>
  </div>

  <h2>Recitals</h2>
  <div class="section">{{ data.recitals }}</div>

  {% if data.definitions %}
  <h2>Definitions</h2>
  <div class="section">
    {% for defn in data.definitions %}
    <p><span class="def-term">"{{ defn.term }}"</span> — {{ defn.definition }}</p>
    {% endfor %}
  </div>
  {% endif %}

  <h2>Scope of Services</h2>
  <div class="section">{{ data.scope_of_services }}</div>

  <h2>Payment Terms</h2>
  <div class="section">{{ data.payment_terms }}</div>

  <h2>Duration</h2>
  <div class="section">{{ data.duration }}</div>

  <h2>Confidentiality</h2>
  <div class="section">{{ data.confidentiality }}</div>

  {% if data.intellectual_property %}
  <h2>Intellectual Property</h2>
  <div class="section">{{ data.intellectual_property }}</div>
  {% endif %}

  <h2>Termination</h2>
  <div class="section">{{ data.termination }}</div>

  <h2>Dispute Resolution</h2>
  <div class="section">{{ data.dispute_resolution }}</div>

  <h2>Force Majeure</h2>
  <div class="section">{{ data.force_majeure }}</div>

  <h2>Governing Law</h2>
  <div class="section">{{ data.governing_law }}</div>

  <h2>Entire Agreement</h2>
  <div class="section">{{ data.entire_agreement }}</div>

  <div class="signature-block">
    <div class="sig-box">
      <div class="sig-line">
        <strong>{{ data.signature_block.party_a.name }}</strong><br>
        {{ data.signature_block.party_a.position }}<br>
        Date: {{ data.signature_block.party_a.date }}
      </div>
      <div class="seal-space">Company Seal</div>
    </div>
    <div class="sig-box">
      <div class="sig-line">
        <strong>{{ data.signature_block.party_b.name }}</strong><br>
        {{ data.signature_block.party_b.position }}<br>
        Date: {{ data.signature_block.party_b.date }}
      </div>
      <div class="seal-space">Company Seal</div>
    </div>
  </div>

  <div class="footer">Generated by <strong>BizPilot NG</strong> | bizpilot.ng</div>
</div>
</body>
</html>
"""

SOCIAL_HTML = """
<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8">
<style>
  body { font-family: 'Segoe UI', Arial, sans-serif; font-size: 11pt; color: #1a1a1a; margin: 0; background: #fafafa; }
  .page { padding: 40px 50px; max-width: 750px; margin: 0 auto; }
  .platform-badge { display: inline-block; background: #E65100; color: white; padding: 6px 16px; border-radius: 20px; font-size: 10pt; font-weight: bold; text-transform: uppercase; letter-spacing: 1px; }
  h1 { color: #E65100; font-size: 18pt; margin: 15px 0 5px; }
  .subtitle { font-size: 10pt; color: #777; margin-bottom: 20px; }
  .card { background: white; border: 1px solid #e0e0e0; border-radius: 10px; padding: 22px; margin-bottom: 18px; box-shadow: 0 1px 3px rgba(0,0,0,0.06); }
  .card-header { display: flex; justify-content: space-between; align-items: center; margin-bottom: 12px; }
  .version-badge { background: #FFF3E0; color: #E65100; padding: 3px 12px; border-radius: 12px; font-size: 9pt; font-weight: bold; }
  .char-count { font-size: 9pt; color: #999; }
  .caption { font-size: 11pt; line-height: 1.7; color: #333; white-space: pre-wrap; }
  .hashtags { margin-top: 12px; }
  .hashtag { display: inline-block; background: #FFF3E0; color: #E65100; padding: 3px 10px; border-radius: 12px; font-size: 9pt; margin: 3px 4px 3px 0; }
  .tip-box { background: #FFF8E1; border-left: 4px solid #FFA000; padding: 14px 18px; margin-top: 20px; border-radius: 0 8px 8px 0; }
  .tip-label { font-size: 9pt; font-weight: bold; color: #F57F17; text-transform: uppercase; letter-spacing: 1px; }
  .tip-text { font-size: 10pt; color: #555; margin-top: 5px; }
  .footer { text-align: center; margin-top: 25px; font-size: 9pt; color: #888; }
  {% if watermark %}.watermark { position: fixed; top: 50%; left: 50%; transform: translate(-50%,-50%) rotate(-45deg); font-size: 60pt; color: rgba(200,200,200,0.25); z-index: -1; white-space: nowrap; }{% endif %}
</style>
</head>
<body>
{% if watermark %}<div class="watermark">BIZPILOT NG FREE</div>{% endif %}
<div class="page">
  <div>
    <span class="platform-badge">{{ data.platform }}</span>
    <h1>Social Media Content</h1>
    <div class="subtitle">3 variations — pick your favourite</div>
  </div>

  {% for v in data.variations %}
  <div class="card">
    <div class="card-header">
      <span class="version-badge">Version {{ v.version }}</span>
      <span class="char-count">{{ v.character_count }} characters</span>
    </div>
    <div class="caption">{{ v.caption }}</div>
    {% if v.hashtags %}
    <div class="hashtags">
      {% for tag in v.hashtags %}<span class="hashtag">{{ tag }}</span>{% endfor %}
    </div>
    {% endif %}
  </div>
  {% endfor %}

  {% if data.posting_tip %}
  <div class="tip-box">
    <div class="tip-label">Posting Tip</div>
    <div class="tip-text">{{ data.posting_tip }}</div>
  </div>
  {% endif %}

  <div class="footer">Generated by <strong>BizPilot NG</strong> | bizpilot.ng</div>
</div>
</body>
</html>
"""

REPLY_HTML = """
<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8">
<style>
  body { font-family: 'Segoe UI', Arial, sans-serif; font-size: 11pt; color: #1a1a1a; margin: 0; }
  .page { padding: 40px 50px; max-width: 750px; margin: 0 auto; }
  .header { border-bottom: 3px solid #00695C; padding-bottom: 18px; margin-bottom: 25px; }
  h1 { color: #00695C; font-size: 18pt; margin: 0 0 8px; }
  .subject { font-size: 12pt; color: #333; }
  .tone-badge { display: inline-block; background: #E0F2F1; color: #00695C; padding: 3px 12px; border-radius: 12px; font-size: 9pt; font-weight: bold; margin-top: 8px; }
  h2 { color: #00695C; font-size: 13pt; margin-top: 25px; border-bottom: 1px solid #b2dfdb; padding-bottom: 5px; }
  .reply-body { background: #f5f5f5; border-left: 4px solid #00695C; padding: 20px 24px; margin: 15px 0; line-height: 1.8; font-size: 10.5pt; white-space: pre-wrap; border-radius: 0 8px 8px 0; }
  .alt-box { background: #E0F2F1; border: 1px solid #b2dfdb; border-radius: 8px; padding: 20px 24px; margin: 15px 0; }
  .alt-label { font-size: 9pt; font-weight: bold; color: #00695C; text-transform: uppercase; letter-spacing: 1px; margin-bottom: 10px; }
  .alt-body { line-height: 1.8; font-size: 10.5pt; white-space: pre-wrap; }
  .point { padding: 4px 0; font-size: 10pt; color: #444; }
  .point:before { content: "\2022"; color: #00695C; font-weight: bold; margin-right: 8px; }
  .action-box { background: #FFF8E1; border-left: 4px solid #FFA000; padding: 14px 18px; margin-top: 20px; border-radius: 0 8px 8px 0; }
  .action-label { font-size: 9pt; font-weight: bold; color: #F57F17; text-transform: uppercase; letter-spacing: 1px; }
  .action-text { font-size: 10pt; color: #555; margin-top: 5px; }
  .footer { text-align: center; margin-top: 30px; font-size: 9pt; color: #888; border-top: 1px solid #e8e8e8; padding-top: 15px; }
  {% if watermark %}.watermark { position: fixed; top: 50%; left: 50%; transform: translate(-50%,-50%) rotate(-45deg); font-size: 60pt; color: rgba(200,200,200,0.25); z-index: -1; white-space: nowrap; }{% endif %}
</style>
</head>
<body>
{% if watermark %}<div class="watermark">BIZPILOT NG FREE</div>{% endif %}
<div class="page">
  <div class="header">
    <h1>Customer Reply</h1>
    <div class="subject"><strong>Subject:</strong> {{ data.subject_line }}</div>
    <span class="tone-badge">{{ data.tone_used }}</span>
  </div>

  <h2>Reply</h2>
  <div class="reply-body">{{ data.reply_text }}</div>

  {% if data.key_points_addressed %}
  <h2>Key Points Addressed</h2>
  <div>
    {% for point in data.key_points_addressed %}
    <div class="point">{{ point }}</div>
    {% endfor %}
  </div>
  {% endif %}

  {% if data.alternative_version %}
  <h2>Alternative Version</h2>
  <div class="alt-box">
    <div class="alt-label">Alternative Wording</div>
    <div class="alt-body">{{ data.alternative_version }}</div>
  </div>
  {% endif %}

  {% if data.follow_up_action %}
  <div class="action-box">
    <div class="action-label">Recommended Follow-up</div>
    <div class="action-text">{{ data.follow_up_action }}</div>
  </div>
  {% endif %}

  <div class="footer">Generated by <strong>BizPilot NG</strong> | bizpilot.ng</div>
</div>
</body>
</html>
"""

BIZPLAN_HTML = """
<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8">
<style>
  body { font-family: Georgia, serif; font-size: 10.5pt; color: #1a1a1a; margin: 0; }
  .page { padding: 45px 55px; max-width: 750px; margin: 0 auto; }
  .header { border-bottom: 3px solid #0D47A1; padding-bottom: 20px; margin-bottom: 25px; text-align: center; }
  h1 { color: #0D47A1; font-size: 20pt; margin: 0 0 5px; }
  .biz-name { font-size: 14pt; color: #333; }
  h2 { color: #0D47A1; font-size: 12pt; border-bottom: 1px solid #bbdefb; padding-bottom: 5px; margin-top: 22px; text-transform: uppercase; letter-spacing: 1px; }
  .section { margin-bottom: 16px; line-height: 1.7; }
  .exec-box { background: #E3F2FD; border-left: 5px solid #0D47A1; padding: 18px 22px; margin: 15px 0; line-height: 1.7; }
  .overview-table { width: 100%; border-collapse: collapse; margin: 10px 0; }
  .overview-table td { padding: 8px 12px; border-bottom: 1px solid #e8e8e8; font-size: 10pt; }
  .overview-table td:first-child { font-weight: bold; color: #0D47A1; width: 35%; }
  .fin-table { width: 100%; border-collapse: collapse; margin: 12px 0; }
  .fin-table th { background: #0D47A1; color: white; padding: 8px 12px; text-align: left; font-size: 10pt; }
  .fin-table td { padding: 8px 12px; border-bottom: 1px solid #e8e8e8; font-size: 10pt; }
  .fin-table tr:nth-child(even) { background: #f5f8ff; }
  .fin-highlight { font-weight: bold; color: #0D47A1; }
  .funding-box { background: #FFF3E0; border: 2px solid #E65100; border-radius: 8px; padding: 20px; margin: 20px 0; }
  .funding-title { font-size: 13pt; font-weight: bold; color: #E65100; margin-bottom: 10px; }
  .funding-amount { font-size: 18pt; font-weight: bold; color: #E65100; }
  .funding-detail { font-size: 10pt; color: #555; margin-top: 8px; line-height: 1.6; }
  .risk-list { list-style: none; padding: 0; }
  .risk-list li { padding: 6px 0 6px 20px; position: relative; font-size: 10pt; }
  .risk-list li:before { content: "\26A0"; position: absolute; left: 0; }
  .footer { text-align: center; margin-top: 30px; font-size: 9pt; color: #888; border-top: 1px solid #e8e8e8; padding-top: 15px; }
  {% if watermark %}.watermark { position: fixed; top: 50%; left: 50%; transform: translate(-50%,-50%) rotate(-45deg); font-size: 60pt; color: rgba(200,200,200,0.25); z-index: -1; white-space: nowrap; }{% endif %}
</style>
</head>
<body>
{% if watermark %}<div class="watermark">BIZPILOT NG FREE</div>{% endif %}
<div class="page">
  <div class="header">
    <h1>Business Plan Summary</h1>
    <div class="biz-name">{{ data.business_overview.name }}</div>
  </div>

  <h2>Executive Summary</h2>
  <div class="exec-box">{{ data.executive_summary }}</div>

  <h2>Business Overview</h2>
  <table class="overview-table">
    <tr><td>Business Name</td><td>{{ data.business_overview.name }}</td></tr>
    <tr><td>Industry</td><td>{{ data.business_overview.industry }}</td></tr>
    <tr><td>Business Type</td><td>{{ data.business_overview.business_type }}</td></tr>
    {% if data.business_overview.year_established %}<tr><td>Year Established</td><td>{{ data.business_overview.year_established }}</td></tr>{% endif %}
    {% if data.business_overview.cac_registration %}<tr><td>CAC Registration</td><td>{{ data.business_overview.cac_registration }}</td></tr>{% endif %}
  </table>

  <h2>Products & Services</h2>
  <div class="section">{{ data.products_and_services }}</div>

  <h2>Market Analysis</h2>
  <div class="section">{{ data.market_analysis }}</div>

  <h2>Target Customer</h2>
  <div class="section">{{ data.target_customer }}</div>

  <h2>Competitive Advantage</h2>
  <div class="section">{{ data.competitive_advantage }}</div>

  <h2>Revenue Model</h2>
  <div class="section">{{ data.revenue_model }}</div>

  <h2>Financial Projections</h2>
  <table class="fin-table">
    <tr><th>Metric</th><th>Amount</th></tr>
    <tr><td>Monthly Revenue</td><td class="fin-highlight">₦{{ "{:,.0f}".format(data.financial_projections.monthly_revenue) }}</td></tr>
    <tr><td>Annual Revenue</td><td class="fin-highlight">₦{{ "{:,.0f}".format(data.financial_projections.annual_revenue) }}</td></tr>
    <tr><td>Monthly Expenses</td><td>₦{{ "{:,.0f}".format(data.financial_projections.monthly_expenses) }}</td></tr>
    <tr><td>Monthly Profit</td><td class="fin-highlight">₦{{ "{:,.0f}".format(data.financial_projections.monthly_profit) }}</td></tr>
    <tr><td>Break-even Period</td><td>{{ data.financial_projections.break_even_months }} months</td></tr>
  </table>

  <div class="funding-box">
    <div class="funding-title">Funding Request</div>
    <div class="funding-amount">₦{{ "{:,.0f}".format(data.funding_request.amount) }}</div>
    <div class="funding-detail">
      <strong>Purpose:</strong> {{ data.funding_request.purpose }}<br>
      <strong>Repayment Plan:</strong> {{ data.funding_request.repayment_plan }}
    </div>
  </div>

  {% if data.risk_factors %}
  <h2>Risk Factors</h2>
  <ul class="risk-list">
    {% for risk in data.risk_factors %}<li>{{ risk }}</li>{% endfor %}
  </ul>
  {% endif %}

  {% if data.management_team %}
  <h2>Management Team</h2>
  <div class="section">{{ data.management_team }}</div>
  {% endif %}

  <div class="section" style="margin-top:20px;font-style:italic;color:#555;">{{ data.conclusion }}</div>

  <div class="footer">Generated by <strong>BizPilot NG</strong> | bizpilot.ng</div>
</div>
</body>
</html>
"""

PLAIN_HTML = """
<!DOCTYPE html><html><head><meta charset="UTF-8">
<style>body{font-family:Arial,sans-serif;font-size:11pt;padding:30px;max-width:700px;margin:0 auto;}
pre{white-space:pre-wrap;font-family:inherit;}</style></head>
<body><pre>{{ content }}</pre>
{% if watermark %}<hr><p style="color:#aaa;text-align:center;font-size:9pt">Generated by BizPilot NG (Free Tier)</p>{% endif %}
</body></html>
"""


# ── Main Generator Class ──────────────────────────────────────────────────────

class DocumentGenerator:
    """
    Converts AI-generated JSON data into formatted documents.
    Entry point: DocumentGenerator.generate()
    """

    HTML_TEMPLATES = {
        DocType.INVOICE:       INVOICE_HTML,
        DocType.PROPOSAL:      PROPOSAL_HTML,
        DocType.CONTRACT:      CONTRACT_HTML,
        DocType.SOCIAL_POST:   SOCIAL_HTML,
        DocType.REPLY:         REPLY_HTML,
        DocType.BUSINESS_PLAN: BIZPLAN_HTML,
    }

    @classmethod
    async def generate(
        cls,
        doc_type: DocType,
        ai_data: dict,
        output_format: OutputFormat,
        subscription_tier: SubscriptionTier = SubscriptionTier.FREE,
    ) -> Optional[bytes]:
        """
        Main entry point.

        Returns raw bytes of the generated file,
        or None if generation fails.
        """
        watermark = subscription_tier == SubscriptionTier.FREE

        if output_format == OutputFormat.TEXT:
            return cls._to_plain_text(doc_type, ai_data).encode("utf-8")

        if output_format == OutputFormat.PDF:
            return cls._to_pdf(doc_type, ai_data, watermark)

        if output_format == OutputFormat.DOCX:
            return cls._to_docx(doc_type, ai_data, watermark)

        return None

    @classmethod
    def _to_plain_text(cls, doc_type: DocType, data: dict) -> str:
        """
        Format the AI JSON as clean plain text for Telegram in-chat display.
        Used for Free tier users.
        """
        if doc_type == DocType.INVOICE:
            return cls._invoice_to_text(data)
        elif doc_type == DocType.PROPOSAL:
            return cls._proposal_to_text(data)
        elif doc_type == DocType.CONTRACT:
            return cls._contract_to_text(data)
        elif doc_type == DocType.SOCIAL_POST:
            return cls._social_to_text(data)
        elif doc_type == DocType.REPLY:
            return cls._reply_to_text(data)
        elif doc_type == DocType.BUSINESS_PLAN:
            return cls._bizplan_to_text(data)
        return str(data)

    @classmethod
    def _to_pdf(cls, doc_type: DocType, data: dict, watermark: bool) -> Optional[bytes]:
        """Render HTML template to PDF using WeasyPrint."""
        if not WEASYPRINT_AVAILABLE:
            logger.warning("WeasyPrint unavailable — falling back to plain text")
            return None

        template_str = cls.HTML_TEMPLATES.get(doc_type, PLAIN_HTML)
        env = Environment(loader=DictLoader({"doc": template_str}))
        template = env.get_template("doc")

        html_content = template.render(data=data, watermark=watermark)

        try:
            pdf_bytes = WeasyHTML(string=html_content).write_pdf()
            return pdf_bytes
        except Exception as e:
            logger.error(f"WeasyPrint PDF error: {e}")
            return None

    @classmethod
    def _to_docx(cls, doc_type: DocType, data: dict, watermark: bool) -> Optional[bytes]:
        """Generate a DOCX document using python-docx."""
        if not DOCX_AVAILABLE:
            logger.warning("python-docx unavailable")
            return None

        try:
            doc = DocxDocument()
            for section in doc.sections:
                section.top_margin    = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin   = Inches(1.2)
                section.right_margin  = Inches(1.2)

            builders = {
                DocType.INVOICE:       cls._build_invoice_docx,
                DocType.PROPOSAL:      cls._build_proposal_docx,
                DocType.CONTRACT:      cls._build_contract_docx,
                DocType.SOCIAL_POST:   cls._build_social_docx,
                DocType.REPLY:         cls._build_reply_docx,
                DocType.BUSINESS_PLAN: cls._build_bizplan_docx,
            }
            builder = builders.get(doc_type)
            if builder:
                builder(doc, data)
            else:
                plain = cls._to_plain_text(doc_type, data)
                for line in plain.split("\n"):
                    doc.add_paragraph(line)

            if watermark:
                doc.add_paragraph("\n[Generated by BizPilot NG — Free Tier]")

            buf = io.BytesIO()
            doc.save(buf)
            return buf.getvalue()
        except Exception as e:
            logger.error(f"DOCX generation error: {e}")
            return None

    # ── Plain Text Formatters ────────────────────────────────────────────────

    @staticmethod
    def _invoice_to_text(d: dict) -> str:
        lines = [
            "=" * 50,
            f"  INVOICE — {d.get('invoice_number', '')}",
            "=" * 50,
            f"Issue Date : {d.get('issue_date', '')}",
            f"Due Date   : {d.get('due_date', '')}",
            "",
            f"FROM: {d.get('seller', {}).get('name', '')}",
            f"TO  : {d.get('bill_to', {}).get('name', '')}",
            f"      {d.get('bill_to', {}).get('email', '')}",
            "",
            "-" * 50,
            f"{'DESCRIPTION':<28} {'QTY':>4} {'UNIT':>8} {'TOTAL':>8}",
            "-" * 50,
        ]
        for item in d.get("items", []):
            lines.append(
                f"{item.get('description', '')[:27]:<28} "
                f"{item.get('quantity', 1):>4} "
                f"{item.get('unit_price', 0):>8,.0f} "
                f"{item.get('line_total', 0):>8,.0f}"
            )
        lines += [
            "-" * 50,
            f"{'Subtotal':<38} ₦{d.get('subtotal', 0):>9,.2f}",
            f"{'VAT (' + str(d.get('vat_rate', 7.5)) + '%)':<38} ₦{d.get('vat_amount', 0):>9,.2f}",
        ]
        if d.get("wht_amount", 0) > 0:
            lines.append(f"{'WHT deduction':<38}-₦{d.get('wht_amount', 0):>9,.2f}")
        lines += [
            "=" * 50,
            f"{'TOTAL DUE':<38} ₦{d.get('total_payable', 0):>9,.2f}",
            "=" * 50,
            "",
            "PAYMENT DETAILS:",
            f"  Account Name   : {d.get('bank_details', {}).get('account_name', '')}",
            f"  Account Number : {d.get('bank_details', {}).get('account_number', '')}",
            f"  Bank           : {d.get('bank_details', {}).get('bank_name', '')}",
            f"  Terms          : {d.get('payment_terms', '')}",
            "",
            d.get("footer", "Thank you for your business."),
        ]
        return "\n".join(lines)

    @staticmethod
    def _proposal_to_text(d: dict) -> str:
        lines = [
            "=" * 50,
            f"  BUSINESS PROPOSAL — {d.get('proposal_number', '')}",
            "=" * 50,
            f"Date: {d.get('date', '')}  |  Valid Until: {d.get('valid_until', '')}",
            f"Prepared for: {d.get('to', {}).get('name', '')}, {d.get('to', {}).get('company', '')}",
            f"Prepared by : {d.get('from', {}).get('name', '')}, {d.get('from', {}).get('business', '')}",
            "",
            "EXECUTIVE SUMMARY",
            "-" * 40,
            d.get("executive_summary", ""),
            "",
            "SCOPE OF WORK",
            "-" * 40,
        ]
        scope = d.get("scope_of_work", "")
        if isinstance(scope, list):
            for s in scope:
                lines.append(f"  - {s}")
        else:
            lines.append(scope)
        lines += [
            "",
            "DELIVERABLES",
            "-" * 40,
        ]
        for i, item in enumerate(d.get("deliverables", []), 1):
            if isinstance(item, dict):
                label = item.get("item", str(item))
                timeline = item.get("timeline", "")
                lines.append(f"  {i}. {label}" + (f" — {timeline}" if timeline else ""))
            else:
                lines.append(f"  {i}. {item}")
        inv = d.get("investment", {})
        lines += [
            "",
            "INVESTMENT",
            "-" * 40,
            f"  Total: ₦{inv.get('total_amount', 0):,.2f}",
            f"  Terms: {inv.get('payment_terms', '')}",
            "",
            d.get("closing_message", ""),
        ]
        return "\n".join(lines)

    @staticmethod
    def _contract_to_text(d: dict) -> str:
        parties = d.get("parties", {})
        pa = parties.get("party_a", {})
        pb = parties.get("party_b", {})
        lines = [
            "=" * 50,
            f"  {d.get('contract_title', 'CONTRACT').upper()}",
            f"  Ref: {d.get('contract_number', '')}  |  Date: {d.get('date', '')}",
            "=" * 50,
            f"PARTY A: {pa.get('name', '')} ({pa.get('role', '')})",
            f"         {pa.get('address', '')}",
            f"PARTY B: {pb.get('name', '')} ({pb.get('role', '')})",
            f"         {pb.get('address', '')}",
            "",
            "RECITALS", "-"*40, d.get("recitals", ""),
            "",
            "SCOPE OF SERVICES", "-"*40, d.get("scope_of_services", ""),
            "",
            "PAYMENT TERMS", "-"*40, d.get("payment_terms", ""),
            "",
            "DURATION", "-"*40, d.get("duration", ""),
            "",
            "CONFIDENTIALITY", "-"*40, d.get("confidentiality", ""),
            "",
            "TERMINATION", "-"*40, d.get("termination", ""),
            "",
            "GOVERNING LAW", "-"*40, d.get("governing_law", ""),
            "",
            "=" * 50,
            "SIGNATURES",
            "",
            f"Party A: ________________  Date: ________________",
            f"({pa.get('name', '')})",
            "",
            f"Party B: ________________  Date: ________________",
            f"({pb.get('name', '')})",
        ]
        return "\n".join(lines)

    @staticmethod
    def _social_to_text(d: dict) -> str:
        lines = [
            f"SOCIAL MEDIA CONTENT — {d.get('platform', '').upper()}",
            "=" * 50,
        ]
        for v in d.get("variations", []):
            lines += [
                f"\n-- Version {v.get('version', '')} --",
                v.get("caption", ""),
                f"Hashtags: {' '.join(v.get('hashtags', []))}",
                f"({v.get('character_count', 0)} chars)",
            ]
        if d.get("posting_tip"):
            lines += ["", f"Tip: {d.get('posting_tip')}"]
        return "\n".join(lines)

    @staticmethod
    def _reply_to_text(d: dict) -> str:
        return "\n".join([
            "CUSTOMER REPLY",
            "=" * 50,
            f"Subject: {d.get('subject_line', '')}",
            "",
            d.get("reply_text", ""),
            "",
            "-- Alternative version --",
            d.get("alternative_version", ""),
            "",
            f"Next step: {d.get('follow_up_action', '')}",
        ])

    @staticmethod
    def _bizplan_to_text(d: dict) -> str:
        fin = d.get("financial_projections", {})
        fund = d.get("funding_request", {})
        lines = [
            "=" * 50,
            "  BUSINESS PLAN SUMMARY",
            "=" * 50,
            "",
            "EXECUTIVE SUMMARY", "-"*40, d.get("executive_summary", ""),
            "",
            "BUSINESS OVERVIEW", "-"*40,
            f"  Name     : {d.get('business_overview', {}).get('name', '')}",
            f"  Industry : {d.get('business_overview', {}).get('industry', '')}",
            f"  CAC Reg  : {d.get('business_overview', {}).get('cac_registration', '')}",
            "",
            "MARKET ANALYSIS", "-"*40, d.get("market_analysis", ""),
            "",
            "COMPETITIVE ADVANTAGE", "-"*40, d.get("competitive_advantage", ""),
            "",
            "FINANCIAL PROJECTIONS", "-"*40,
            f"  Monthly Revenue  : ₦{fin.get('monthly_revenue', 0):,.0f}",
            f"  Annual Revenue   : ₦{fin.get('annual_revenue', 0):,.0f}",
            f"  Monthly Expenses : ₦{fin.get('monthly_expenses', 0):,.0f}",
            f"  Monthly Profit   : ₦{fin.get('monthly_profit', 0):,.0f}",
            f"  Break-even       : {fin.get('break_even_months', 0)} months",
            "",
            "FUNDING REQUEST", "-"*40,
            f"  Amount  : ₦{fund.get('amount', 0):,.0f}",
            f"  Purpose : {fund.get('purpose', '')}",
            f"  Repayment: {fund.get('repayment_plan', '')}",
            "",
            d.get("conclusion", ""),
        ]
        return "\n".join(lines)

    # ── DOCX Builders ────────────────────────────────────────────────────────

    @staticmethod
    def _build_invoice_docx(doc: "DocxDocument", d: dict) -> None:
        """Populate a python-docx Document with invoice content."""
        GREEN = RGBColor(0x1B, 0x5E, 0x20)

        title = doc.add_heading("INVOICE", 0)
        title.runs[0].font.color.rgb = GREEN

        doc.add_paragraph(f"Invoice No: {d.get('invoice_number', '')}")
        doc.add_paragraph(f"Date: {d.get('issue_date', '')}  |  Due: {d.get('due_date', '')}")
        doc.add_paragraph(f"From: {d.get('seller', {}).get('name', '')}")
        doc.add_paragraph(f"To:   {d.get('bill_to', {}).get('name', '')}")
        doc.add_paragraph("")

        table = doc.add_table(rows=1, cols=4)
        table.style = "Light List Accent 1"
        hdr = table.rows[0].cells
        for i, h in enumerate(["Description", "Qty", "Unit Price", "Total"]):
            hdr[i].text = h

        for item in d.get("items", []):
            row = table.add_row().cells
            row[0].text = item.get("description", "")
            row[1].text = str(item.get("quantity", 1))
            row[2].text = f"₦{item.get('unit_price', 0):,.2f}"
            row[3].text = f"₦{item.get('line_total', 0):,.2f}"

        doc.add_paragraph("")
        doc.add_paragraph(f"Subtotal: ₦{d.get('subtotal', 0):,.2f}")
        doc.add_paragraph(f"VAT ({d.get('vat_rate', 7.5)}%): ₦{d.get('vat_amount', 0):,.2f}")
        total_para = doc.add_paragraph(f"TOTAL DUE: ₦{d.get('total_payable', 0):,.2f}")
        total_para.runs[0].bold = True

        doc.add_paragraph("")
        bd = d.get("bank_details", {})
        doc.add_paragraph(f"Pay to: {bd.get('account_name', '')} | {bd.get('account_number', '')} | {bd.get('bank_name', '')}")

    @staticmethod
    def _build_proposal_docx(doc: "DocxDocument", d: dict) -> None:
        NAVY = RGBColor(0x1A, 0x23, 0x7E)

        title = doc.add_heading("BUSINESS PROPOSAL", 0)
        title.runs[0].font.color.rgb = NAVY

        meta = doc.add_paragraph()
        meta.add_run(f"Ref: {d.get('proposal_number', '')}  |  Date: {d.get('date', '')}  |  Valid until: {d.get('valid_until', '')}").font.size = Pt(9)
        doc.add_paragraph(f"Prepared for: {d.get('to', {}).get('name', '')}, {d.get('to', {}).get('company', '')}")
        doc.add_paragraph(f"Prepared by: {d.get('from', {}).get('name', '')}, {d.get('from', {}).get('business', '')}")

        h = doc.add_heading("Executive Summary", level=1)
        h.runs[0].font.color.rgb = NAVY
        doc.add_paragraph(d.get("executive_summary", ""))

        h = doc.add_heading("Project Overview", level=1)
        h.runs[0].font.color.rgb = NAVY
        doc.add_paragraph(d.get("project_overview", ""))

        h = doc.add_heading("Scope of Work", level=1)
        h.runs[0].font.color.rgb = NAVY
        scope = d.get("scope_of_work", "")
        if isinstance(scope, list):
            for s in scope:
                doc.add_paragraph(s, style="List Bullet")
        else:
            doc.add_paragraph(scope)

        h = doc.add_heading("Deliverables", level=1)
        h.runs[0].font.color.rgb = NAVY
        for i, item in enumerate(d.get("deliverables", []), 1):
            if isinstance(item, dict):
                label = item.get("item", str(item))
                timeline = item.get("timeline", "")
                doc.add_paragraph(f"{i}. {label}" + (f" — {timeline}" if timeline else ""))
            else:
                doc.add_paragraph(f"{i}. {item}")

        h = doc.add_heading("Timeline", level=1)
        h.runs[0].font.color.rgb = NAVY
        doc.add_paragraph(d.get("timeline", ""))

        h = doc.add_heading("Investment", level=1)
        h.runs[0].font.color.rgb = NAVY
        inv = d.get("investment", {})
        total_para = doc.add_paragraph()
        total_run = total_para.add_run(f"Total Project Value: ₦{inv.get('total_amount', 0):,.2f}")
        total_run.bold = True
        total_run.font.size = Pt(14)
        total_run.font.color.rgb = NAVY

        breakdown = inv.get("breakdown", [])
        if breakdown:
            table = doc.add_table(rows=1, cols=3)
            table.style = "Light List Accent 1"
            hdr = table.rows[0].cells
            for i, h_text in enumerate(["Milestone", "Percentage", "Amount"]):
                hdr[i].text = h_text
            for m in breakdown:
                row = table.add_row().cells
                row[0].text = m.get("milestone", "")
                row[1].text = f"{m.get('percentage', 0)}%"
                row[2].text = f"₦{m.get('amount', 0):,.2f}"

        doc.add_paragraph(f"Payment Terms: {inv.get('payment_terms', '')}")

        h = doc.add_heading("Why Choose Us", level=1)
        h.runs[0].font.color.rgb = NAVY
        doc.add_paragraph(d.get("why_choose_us", ""))

        h = doc.add_heading("Terms & Conditions", level=1)
        h.runs[0].font.color.rgb = NAVY
        doc.add_paragraph(d.get("terms_and_conditions", ""))

        doc.add_paragraph(d.get("closing_message", ""))

        doc.add_paragraph("")
        sig_table = doc.add_table(rows=2, cols=2)
        sig_table.cell(0, 0).text = f"{d.get('from', {}).get('name', '')}\n{d.get('from', {}).get('business', '')}\nDate: ________________"
        sig_table.cell(0, 1).text = f"{d.get('to', {}).get('name', '')}\n{d.get('to', {}).get('company', '')}\nDate: ________________"
        sig_table.cell(1, 0).text = "Signature: ________________"
        sig_table.cell(1, 1).text = "Signature: ________________"

    @staticmethod
    def _build_contract_docx(doc: "DocxDocument", d: dict) -> None:
        CHARCOAL = RGBColor(0x26, 0x32, 0x38)

        title = doc.add_heading(d.get("contract_title", "CONTRACT").upper(), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.color.rgb = CHARCOAL

        ref_para = doc.add_paragraph()
        ref_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ref_run = ref_para.add_run(f"Ref: {d.get('contract_number', '')}  |  Date: {d.get('date', '')}")
        ref_run.font.size = Pt(9)
        ref_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        parties = d.get("parties", {})
        pa = parties.get("party_a", {})
        pb = parties.get("party_b", {})

        doc.add_paragraph("")
        p_table = doc.add_table(rows=1, cols=2)
        cell_a = p_table.cell(0, 0)
        cell_a.text = f"PARTY A — {pa.get('role', '')}\n{pa.get('name', '')}\n{pa.get('address', '')}"
        cell_b = p_table.cell(0, 1)
        cell_b.text = f"PARTY B — {pb.get('role', '')}\n{pb.get('name', '')}\n{pb.get('address', '')}"

        sections = [
            ("Recitals", "recitals"),
            ("Scope of Services", "scope_of_services"),
            ("Payment Terms", "payment_terms"),
            ("Duration", "duration"),
            ("Confidentiality", "confidentiality"),
            ("Intellectual Property", "intellectual_property"),
            ("Termination", "termination"),
            ("Dispute Resolution", "dispute_resolution"),
            ("Force Majeure", "force_majeure"),
            ("Governing Law", "governing_law"),
            ("Entire Agreement", "entire_agreement"),
        ]

        for heading_text, key in sections:
            content = d.get(key, "")
            if not content:
                continue
            h = doc.add_heading(heading_text.upper(), level=1)
            h.runs[0].font.color.rgb = CHARCOAL
            doc.add_paragraph(content)

        definitions = d.get("definitions", [])
        if definitions:
            h = doc.add_heading("DEFINITIONS", level=1)
            h.runs[0].font.color.rgb = CHARCOAL
            for defn in definitions:
                p = doc.add_paragraph()
                bold_run = p.add_run(f'"{defn.get("term", "")}"')
                bold_run.bold = True
                p.add_run(f' — {defn.get("definition", "")}')

        doc.add_paragraph("")
        doc.add_paragraph("")
        sig_block = d.get("signature_block", {})
        sig_a = sig_block.get("party_a", {})
        sig_b = sig_block.get("party_b", {})

        sig_table = doc.add_table(rows=3, cols=2)
        sig_table.cell(0, 0).text = "PARTY A"
        sig_table.cell(0, 1).text = "PARTY B"
        sig_table.cell(1, 0).text = f"\n\nSignature: ________________\n{sig_a.get('name', '')}\n{sig_a.get('position', '')}\nDate: {sig_a.get('date', '')}"
        sig_table.cell(1, 1).text = f"\n\nSignature: ________________\n{sig_b.get('name', '')}\n{sig_b.get('position', '')}\nDate: {sig_b.get('date', '')}"
        sig_table.cell(2, 0).text = "[Company Seal]"
        sig_table.cell(2, 1).text = "[Company Seal]"

    @staticmethod
    def _build_social_docx(doc: "DocxDocument", d: dict) -> None:
        ORANGE = RGBColor(0xE6, 0x51, 0x00)

        title = doc.add_heading("SOCIAL MEDIA CONTENT", 0)
        title.runs[0].font.color.rgb = ORANGE

        platform_para = doc.add_paragraph()
        run = platform_para.add_run(f"Platform: {d.get('platform', '').upper()}")
        run.bold = True
        run.font.color.rgb = ORANGE
        doc.add_paragraph("3 variations — pick your favourite")
        doc.add_paragraph("")

        for v in d.get("variations", []):
            h = doc.add_heading(f"Version {v.get('version', '')}", level=2)
            h.runs[0].font.color.rgb = ORANGE

            caption_para = doc.add_paragraph(v.get("caption", ""))
            caption_para.paragraph_format.space_after = Pt(6)

            hashtags = v.get("hashtags", [])
            if hashtags:
                tag_para = doc.add_paragraph()
                tag_run = tag_para.add_run(" ".join(hashtags))
                tag_run.font.size = Pt(9)
                tag_run.font.color.rgb = ORANGE

            count_para = doc.add_paragraph()
            count_run = count_para.add_run(f"{v.get('character_count', 0)} characters")
            count_run.font.size = Pt(8)
            count_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

            doc.add_paragraph("")

        tip = d.get("posting_tip", "")
        if tip:
            doc.add_paragraph("")
            tip_para = doc.add_paragraph()
            tip_label = tip_para.add_run("POSTING TIP: ")
            tip_label.bold = True
            tip_label.font.color.rgb = RGBColor(0xF5, 0x7F, 0x17)
            tip_para.add_run(tip)

    @staticmethod
    def _build_reply_docx(doc: "DocxDocument", d: dict) -> None:
        TEAL = RGBColor(0x00, 0x69, 0x5C)

        title = doc.add_heading("CUSTOMER REPLY", 0)
        title.runs[0].font.color.rgb = TEAL

        subj_para = doc.add_paragraph()
        subj_para.add_run("Subject: ").bold = True
        subj_para.add_run(d.get("subject_line", ""))

        tone_para = doc.add_paragraph()
        tone_run = tone_para.add_run(f"Tone: {d.get('tone_used', '')}")
        tone_run.font.size = Pt(9)
        tone_run.font.color.rgb = TEAL

        h = doc.add_heading("Reply", level=1)
        h.runs[0].font.color.rgb = TEAL
        doc.add_paragraph(d.get("reply_text", ""))

        points = d.get("key_points_addressed", [])
        if points:
            h = doc.add_heading("Key Points Addressed", level=1)
            h.runs[0].font.color.rgb = TEAL
            for point in points:
                p = doc.add_paragraph(point, style="List Bullet")

        alt = d.get("alternative_version", "")
        if alt:
            h = doc.add_heading("Alternative Version", level=1)
            h.runs[0].font.color.rgb = TEAL
            doc.add_paragraph(alt)

        followup = d.get("follow_up_action", "")
        if followup:
            doc.add_paragraph("")
            fp = doc.add_paragraph()
            fp.add_run("Recommended Follow-up: ").bold = True
            fp.add_run(followup)

    @staticmethod
    def _build_bizplan_docx(doc: "DocxDocument", d: dict) -> None:
        BLUE = RGBColor(0x0D, 0x47, 0xA1)

        title = doc.add_heading("BUSINESS PLAN SUMMARY", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.color.rgb = BLUE

        biz = d.get("business_overview", {})
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.add_run(biz.get("name", ""))
        name_run.font.size = Pt(14)
        name_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        h = doc.add_heading("Executive Summary", level=1)
        h.runs[0].font.color.rgb = BLUE
        exec_para = doc.add_paragraph(d.get("executive_summary", ""))
        exec_para.paragraph_format.space_before = Pt(6)

        h = doc.add_heading("Business Overview", level=1)
        h.runs[0].font.color.rgb = BLUE
        overview_rows = [
            ("Business Name", biz.get("name", "")),
            ("Industry", biz.get("industry", "")),
            ("Business Type", biz.get("business_type", "")),
        ]
        if biz.get("year_established"):
            overview_rows.append(("Year Established", str(biz["year_established"])))
        if biz.get("cac_registration"):
            overview_rows.append(("CAC Registration", biz["cac_registration"]))

        table = doc.add_table(rows=len(overview_rows), cols=2)
        table.style = "Light List Accent 1"
        for i, (label, value) in enumerate(overview_rows):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = value

        content_sections = [
            ("Products & Services", "products_and_services"),
            ("Market Analysis", "market_analysis"),
            ("Target Customer", "target_customer"),
            ("Competitive Advantage", "competitive_advantage"),
            ("Revenue Model", "revenue_model"),
        ]
        for heading_text, key in content_sections:
            content = d.get(key, "")
            if not content:
                continue
            h = doc.add_heading(heading_text, level=1)
            h.runs[0].font.color.rgb = BLUE
            doc.add_paragraph(content)

        h = doc.add_heading("Financial Projections", level=1)
        h.runs[0].font.color.rgb = BLUE
        fin = d.get("financial_projections", {})
        fin_data = [
            ("Monthly Revenue", f"₦{fin.get('monthly_revenue', 0):,.0f}"),
            ("Annual Revenue", f"₦{fin.get('annual_revenue', 0):,.0f}"),
            ("Monthly Expenses", f"₦{fin.get('monthly_expenses', 0):,.0f}"),
            ("Monthly Profit", f"₦{fin.get('monthly_profit', 0):,.0f}"),
            ("Break-even Period", f"{fin.get('break_even_months', 0)} months"),
        ]
        fin_table = doc.add_table(rows=len(fin_data) + 1, cols=2)
        fin_table.style = "Light List Accent 1"
        fin_table.cell(0, 0).text = "Metric"
        fin_table.cell(0, 1).text = "Amount"
        for i, (label, value) in enumerate(fin_data, 1):
            fin_table.cell(i, 0).text = label
            cell = fin_table.cell(i, 1)
            cell.text = value

        h = doc.add_heading("Funding Request", level=1)
        h.runs[0].font.color.rgb = BLUE
        fund = d.get("funding_request", {})
        amount_para = doc.add_paragraph()
        amount_run = amount_para.add_run(f"₦{fund.get('amount', 0):,.0f}")
        amount_run.bold = True
        amount_run.font.size = Pt(16)
        amount_run.font.color.rgb = RGBColor(0xE6, 0x51, 0x00)
        doc.add_paragraph(f"Purpose: {fund.get('purpose', '')}")
        doc.add_paragraph(f"Repayment Plan: {fund.get('repayment_plan', '')}")

        risks = d.get("risk_factors", [])
        if risks:
            h = doc.add_heading("Risk Factors", level=1)
            h.runs[0].font.color.rgb = BLUE
            for risk in risks:
                doc.add_paragraph(risk, style="List Bullet")

        mgmt = d.get("management_team", "")
        if mgmt:
            h = doc.add_heading("Management Team", level=1)
            h.runs[0].font.color.rgb = BLUE
            doc.add_paragraph(mgmt)

        conclusion = d.get("conclusion", "")
        if conclusion:
            doc.add_paragraph("")
            conc_para = doc.add_paragraph()
            conc_run = conc_para.add_run(conclusion)
            conc_run.italic = True
